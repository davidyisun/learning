#!/usr/bin/env python
# -*- coding:utf-8 -*-
"""
    脚本名: docx test
Created on 2018-11-07
@author:David Yisun
@group:data
"""
import docx
from docx import Document
doc = Document('./data/123.docx')
import pandas as pd

p = []
for paragragh in doc.paragraphs:
    text = paragragh.text
    if text != '':
        p.append(paragragh.text)

tables = [table for table in doc.tables]

t = []
for table in tables:
    _t = []
    for row in table.rows:
        _t.append([cell.text for cell in row.cells])
    df = pd.DataFrame(_t)
    t.append(df)

# doc 批量转 docx

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table

parent_elm = doc.element.body

o = [i for i in parent_elm.iterchildren()]

l = []
for i in o:
    if isinstance(i, CT_P):
        l.append(Paragraph(i, doc))
    elif isinstance(i, CT_Tbl):
        l.append(Table(i, doc))


paragraph = doc.tables[0]
run = paragraph.add_run()
run.add_picture('./data/123.jpg')
doc.save('./data/234.docx')

# windows




