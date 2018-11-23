#encoding:utf-8
from docx import Document
import os
import sys
import time
import json
import datetime
import platform
import traceback
import subprocess
from werkzeug.utils import secure_filename
from flask import Flask, Blueprint, request, Response
cur_dir = os.path.dirname(os.path.abspath(__file__))
cn_digit_list = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
cn_num_list = []
for i in range(len(cn_digit_list)):
    for j in range(len(cn_digit_list)):
        cn_num_list.append(cn_digit_list[i] + cn_digit_list[j])
cn_num_set = set(cn_digit_list)


def doc2docx(input_file):
    platform_name = platform.platform()
    output_file = input_file + 'x'
    if 'Windows' in platform_name:
        from win32com import client as wc
        word = wc.Dispatch('Word.Application')
        doc = word.Documents.Open(input_file)
        doc.SaveAs(output_file, 16)  # 12 for doc, 16 for docx
        doc.Close()
        word.Quit()
    else:
        cmd = 'soffice --headless --invisible --convert-to docx {} --outdir {}'.format(
            input_file, os.path.dirname(input_file))
        subprocess.call(cmd, shell=True)
    return output_file


def get_tables(input_path, output_path):
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    document = Document(input_path)
    doc_new = Document()
    tables = document.tables
    last_col_text = ''
    pg_num = 1
    for ti in range(len(tables)):
        rows = tables[ti].rows
        print('rows num: {}'.format(len(rows)))
        for tj in range(len(rows)):
            title = '####{}{}{}####'.format('0001', '0002', '%04d' % pg_num)
            pg_num += 1
            doc_new.add_paragraph(title, style='Normal')
            cols = rows[tj].cells
            print('cols num: {}'.format(len(cols)))
            for tk in range(len(cols)):
                col_text = cols[tk].text.strip()
                if col_text != last_col_text:
                    splits = col_text.split('\n')
                    for split in splits:
                        flag = False
                        for cn_num in cn_num_set:
                            if split.startswith(cn_num + '、'):
                                doc_new.add_paragraph('', style='Normal')
                                doc_new.add_paragraph(split, style='Normal')
                                flag = True
                                break
                        if not flag:
                            doc_new.add_paragraph(split, style='Normal')
                    # doc_new.add_paragraph(col_text, style='Normal')
                    last_col_text = col_text
            doc_new.add_paragraph('', style='Normal')
    doc_new.save(output_path)


def batch_process():
    data_dir = os.path.join(cur_dir, '20180903_excel')
    result_dir = data_dir + '_rebuild'
    filenames = sorted(os.listdir(data_dir))
    for filename in filenames:
        t0 = time.time()
        file_path = os.path.join(data_dir, filename)
        if filename.endswith('.doc'):
            file_path = doc2docx(file_path)
        get_tables(file_path, os.path.join(result_dir, os.path.basename(file_path)))
        print(time.time() - t0)


APP = Flask(__name__)
APP.config['JSON_AS_ASCII'] = False
APP.config['ALLOWED_EXTENSIONS'] = {'doc', 'docx'}


@APP.route('/')
def hello_world():
    msg = 'Connect test success! cur_time: {}\n'.format(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    return msg


@APP.route('/rebuild', methods=['POST'])
def parse_docx():
    print('=' * 10 + ' {} /rebuild get request time: {}'.format(
        __file__, datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    t0 = time.time()
    try:
        upload_file = request.files['file']
        filename = secure_filename(upload_file.filename)
        print('filename: {}'.format(filename))
        input_path = os.path.join(data_dir, filename)
        upload_file.save(input_path)
        if input_path.endswith('.doc'):
            input_path = doc2docx(input_path)
        output_path = os.path.join(result_dir, os.path.basename(input_path))
        get_tables(input_path, output_path)
    except Exception as e:
        traceback.print_exc()
        result = {'code': -1, 'message': traceback.format_exc()}
        return json.dumps(result, ensure_ascii=False)
    print('This request takes {} seconds'.format(time.time() - t0))
    # return json.dumps(result, ensure_ascii=False)
    return Response(open(output_path, 'rb'))


@APP.route('/segment', methods=['POST'])
def segment_docx():
    print('=' * 10 + ' {} /segment get request time: {}'.format(
        __file__, datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    t0 = time.time()
    try:
        upload_file = request.files['file']
        filename = secure_filename(upload_file.filename)
        print('filename: {}'.format(filename))
        input_path = os.path.join(data_dir, filename)
        upload_file.save(input_path)
        if input_path.endswith('.doc'):
            input_path = doc2docx(input_path)
        output_path = os.path.join(result_dir, os.path.basename(input_path))
        get_tables(input_path, output_path)
    except Exception as e:
        traceback.print_exc()
        result = {'code': -1, 'message': traceback.format_exc()}
        return json.dumps(result, ensure_ascii=False)
    print('This request takes {} seconds'.format(time.time() - t0))
    # return json.dumps(result, ensure_ascii=False)
    return Response(open(output_path, 'rb'))


if __name__ == '__main__':
    data_dir = os.path.join(cur_dir, 'data')
    result_dir = data_dir + '_rebuild'
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    if not os.path.exists(result_dir):
        os.makedirs(result_dir)
    APP.run(host='0.0.0.0', port=8500)

